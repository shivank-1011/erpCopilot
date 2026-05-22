"""
services/parser.py — Extract text from PDF and DOCX files.

PDF:  Uses PyPDF2, returns list of {page_num, text} dicts.
DOCX: Uses python-docx, groups paragraphs into logical pages.
"""
import hashlib
import io
from pathlib import Path
from typing import BinaryIO

import PyPDF2
from docx import Document as DocxDocument


# ── PDF Parsing ──────────────────────────────────────────────

def extract_pdf(file_bytes: bytes) -> tuple[list[dict], int]:
    """
    Returns:
        pages: list of {"page_num": int, "text": str}
        page_count: total pages
    """
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages = []

    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
            # Sanitize NUL characters (PostgreSQL TEXT columns do not support \x00)
            text = text.replace("\x00", "")
            text = text.strip()
        except Exception:
            text = ""

        if text:
            pages.append({"page_num": i, "text": text})

    if not pages:
        raise ValueError(
            "Could not extract text from this PDF. "
            "It may be a scanned/image-based PDF — OCR support is planned."
        )

    return pages, len(reader.pages)


# ── DOCX Parsing ─────────────────────────────────────────────

def extract_docx(file_bytes: bytes) -> tuple[list[dict], int]:
    """
    DOCX files don't have explicit pages.
    We treat the whole document as page 1 and let the chunker handle splitting.
    Returns:
        pages: [{"page_num": 1, "text": full_text}]
        page_count: 1
    """
    doc = DocxDocument(io.BytesIO(file_bytes))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise ValueError("Could not extract text from this DOCX file.")

    full_text = "\n\n".join(paragraphs)
    return [{"page_num": 1, "text": full_text}], 1


# ── Dispatcher ───────────────────────────────────────────────

def parse_document(
    file_bytes: bytes, filename: str
) -> tuple[list[dict], int, str]:
    """
    Detect file type and dispatch to the correct parser.

    Returns:
        pages:      list of {"page_num": int, "text": str}
        page_count: int
        file_type:  "pdf" | "docx"
    """
    ext = Path(filename).suffix.lower().lstrip(".")

    if ext == "pdf":
        pages, page_count = extract_pdf(file_bytes)
        return pages, page_count, "pdf"
    elif ext in ("docx", "doc"):
        pages, page_count = extract_docx(file_bytes)
        return pages, page_count, "docx"
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Only PDF and DOCX are supported.")


# ── Utility ───────────────────────────────────────────────────

def compute_file_hash(file_bytes: bytes) -> str:
    """SHA-256 hash of raw file bytes — used for duplicate detection."""
    return hashlib.sha256(file_bytes).hexdigest()
